"""
office-agent.py — AI 文档控制引擎 (python-docx/openpyxl + LibreOffice CLI)
端口: 9810 | 支持: .docx/.xlsx/.odt/.ods + LibreOffice 打开显示 + CLI模式
CC维护 | 2026-04-25
"""

import asyncio
import json
import logging
import os
import re
import subprocess
import time
from pathlib import Path
from typing import Optional

import httpx
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
import uvicorn

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger("office-agent")

app = FastAPI(title="Office Agent", version="2.1")
app.add_middleware(
    CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"]
)

LITELLM_URL = "http://localhost:4000/v1"
LITELLM_KEY = "sk-local-8e781a02c87854bf06ed2a5e871915962227ab91bc71937e"
HISTORY_FILE = Path.home() / ".local" / "share" / "office-agent-history.jsonl"
HISTORY_FILE.parent.mkdir(parents=True, exist_ok=True)

# 当前操作的文件（session级别）
_current_file: Optional[str] = None

# ── 意图解析 ─────────────────────────────────────────────

INTENT_SYSTEM = """你是 LibreOffice 文档操作意图解析器。把用户的自然语言指令转换为结构化 JSON。

支持的 action（Writer/.docx）：
- bold / italic / underline: 格式化段落。target="paragraph_N"(从1计) 或 "all" 或 "last"
- font_size: 字体大小。target, size(数字，磅)
- color: 字体颜色。target, color(十六进制如"FF0000"表示红色)
- set_align: 对齐方式。target, align("left"/"center"/"right"/"justify")
- set_line_spacing: 行间距。target, spacing(行距倍数，如1.5)
- insert_text: 在末尾插入文字。content
- insert_table: 插入表格。rows, cols（在末尾）
- find_replace: 查找替换。find, replace
- delete_paragraph: 删除段落。target="paragraph_N"
- set_heading: 设置标题级别。target="paragraph_N", level(1-6)
- add_paragraph: 新增段落。content
- get_content: 获取文档内容（返回全文）
- page_break: 在末尾插入分页符
- insert_image: 插入图片。image_path, width(cm可选), height(cm可选)
- header_footer: 设置页眉页脚。header, footer
- add_comment: 添加批注（需LO运行）。content, target="paragraph_N"
- track_changes: 开/关修订模式（需LO运行）。enable=true/false
- export_pdf: 导出PDF。path(可选，默认同目录)
- save: 保存文档
- open: 用LibreOffice打开文件预览。filepath（可选，默认当前文件）
- new_doc: 新建.docx文档。filepath（保存路径）
- set_current: 设置当前操作文件。filepath

支持的 action（Calc/.xlsx）：
- read_cell: 读取单元格。sheet(可选), cell("A1")
- write_cell: 写入单元格。sheet(可选), cell("A1"), value
- write_row: 写入一行。sheet(可选), row(行号从1), values(列表)
- get_sheet_summary: 获取表格摘要
- read_range: 读取范围。range("A1:C5")
- find_replace: 查找替换。find, replace
- batch_write: 批量写入多单元格。cells=[{"cell":"A1","value":1},{"cell":"B1","value":2}]
- delete_row: 删除行。row(行号)
- insert_row: 插入行。row(行号), values(可选列表)
- format_cell: 格式化单元格。cell, format(数字格式), bold, color(十六进制), size
- auto_fit: 自动调整列宽

输出纯 JSON，不要 markdown 代码块。例：
"把第二段改成加粗" → {"action":"bold","target":"paragraph_2"}
"在末尾插入3行4列的表格" → {"action":"insert_table","rows":3,"cols":4}
"把所有公司替换成企业" → {"action":"find_replace","find":"公司","replace":"企业"}
"新建一个报告文档保存到桌面" → {"action":"new_doc","filepath":"~/Desktop/报告.docx"}
"A1格填入销售额" → {"action":"write_cell","cell":"A1","value":"销售额"}
"读取A1到C5的数据" → {"action":"read_range","range":"A1:C5"}
"删除第3行" → {"action":"delete_row","row":3}
"在第2行插入姓名年龄性别" → {"action":"insert_row","row":2,"values":["姓名","年龄","性别"]}
"批量写入A1=100 B1=200" → {"action":"batch_write","cells":[{"cell":"A1","value":100},{"cell":"B1","value":200}]}
"把A1单元格背景设成黄色加粗" → {"action":"format_cell","cell":"A1","color":"FFFF00","bold":true}
"插入图片logo.png宽度5cm" → {"action":"insert_image","image_path":"~/Desktop/logo.png","width":5}
"设置页眉为公司名称" → {"action":"header_footer","header":"公司名称"}
"打开当前文件" → {"action":"open"}
"保存" → {"action":"save"}
"把标题变成红色" → {"action":"color","target":"paragraph_1","color":"FF0000"}
"居中对齐第一段" → {"action":"set_align","target":"paragraph_1","align":"center"}
"导出为PDF" → {"action":"export_pdf"}
"开启修订模式" → {"action":"track_changes","enable":true}"""


async def parse_intent(text: str) -> dict:
    """用 glm CLI 解析意图（避免 LiteLLM 超时）"""
    prompt = f"{INTENT_SYSTEM}\n\n用户指令：{text}\n\n只输出JSON，不要其他文字："
    proc = await asyncio.create_subprocess_exec(
        "/home/charlie/.local/bin/glm",
        prompt,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.DEVNULL,
    )
    try:
        stdout, _ = await asyncio.wait_for(proc.communicate(), timeout=20)
        raw = stdout.decode().strip()
        # 提取 JSON 块
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            return json.loads(m.group())
        return json.loads(raw)
    except (asyncio.TimeoutError, json.JSONDecodeError, Exception) as e:
        log.warning("GLM解析失败: %s，尝试规则匹配", e)
        return rule_based_intent(text)


def rule_based_intent(text: str) -> dict:
    """规则降级：GLM不可用时的简单意图匹配"""
    t = text.lower()
    if "新建" in t or "创建" in t:
        import re as _re

        # 先匹配"叫xxx.docx"或"名为xxx.docx"模式，提取纯文件名
        m_name = _re.search(
            r"(?:叫|名为|名字叫|文件名|命名为)\s*([\w\u4e00-\u9fff\-_.]+\.docx)", text
        )
        # 再匹配完整路径 ~/Desktop/xxx.docx
        m_path = _re.search(r"(?:~/|/)[^\s]+\.docx", text)
        # "保存到桌面" → ~/Desktop/
        desktop = "桌面" in text or "desktop" in t
        if m_name:
            fname = m_name.group(1)
            fp = f"~/Desktop/{fname}" if desktop else f"~/Desktop/{fname}"
        elif m_path:
            fp = m_path.group()
        else:
            # fallback: 找任意位置的 xxx.docx
            m2 = _re.search(r"([\w\u4e00-\u9fff\-_.]+\.docx)", text)
            fname = m2.group(1) if m2 else "新文档.docx"
            fp = f"~/Desktop/{fname}" if desktop else f"~/Desktop/{fname}"
        return {"action": "new_doc", "filepath": fp}
    if "保存" in t:
        return {"action": "save"}
    if "加粗" in t or "粗体" in t:
        m = re.search(r"第(\d+)段", text)
        return {"action": "bold", "target": f"paragraph_{m.group(1)}" if m else "all"}
    if "斜体" in t:
        return {"action": "italic", "target": "all"}
    if "替换" in t:
        parts = re.findall(r'[「『"](.*?)[」』"]', text)
        if len(parts) >= 2:
            return {"action": "find_replace", "find": parts[0], "replace": parts[1]}
    if "内容" in t or "查看" in t:
        return {"action": "get_content"}
    if "打开" in t:
        return {"action": "open"}
    if "表格" in t:
        rows = re.search(r"(\d+)\s*行", text)
        cols = re.search(r"(\d+)\s*列", text)
        return {
            "action": "insert_table",
            "rows": int(rows.group(1)) if rows else 3,
            "cols": int(cols.group(1)) if cols else 3,
        }
    return {"action": "unknown", "raw": text}


# ── 文档执行引擎 ──────────────────────────────────────────


def resolve_file(filepath: Optional[str] = None) -> Path:
    global _current_file
    fp = filepath or _current_file
    if not fp:
        raise RuntimeError("未设置当前文档，请先发送：打开 /path/to/file.docx")
    p = Path(os.path.expanduser(fp))
    if not p.exists() and not fp.startswith("new:"):
        raise RuntimeError(f"文件不存在: {p}")
    return p


def is_calc(p: Path) -> bool:
    return p.suffix.lower() in (".xlsx", ".xls", ".ods", ".csv")


# ─── Writer 操作 ───────────────────────────────────────


def get_paragraph(doc, target: str):
    import docx

    paras = [p for p in doc.paragraphs if p.text.strip() or p.runs]
    paras_all = doc.paragraphs
    if target == "all":
        return paras_all
    if target == "last":
        return [paras_all[-1]] if paras_all else []
    if target.startswith("paragraph_"):
        idx = int(target.split("_")[1]) - 1
        return [paras_all[idx]] if 0 <= idx < len(paras_all) else []
    return paras_all


def _lo_running() -> bool:
    """检测 LibreOffice GUI 是否正在运行"""
    try:
        r = subprocess.run(["pgrep", "-x", "soffice"], capture_output=True)
        return r.returncode == 0
    except Exception:
        return False


def exec_via_macro(intent: dict, filepath: str) -> dict:
    """通过 LibreOffice 内嵌 Python 宏执行（完整 UNO 访问）"""
    cmd_file = Path.home() / ".local/state/lo_cmd.json"
    result_file = Path.home() / ".local/state/lo_result.json"

    # 写命令
    cmd_data = dict(intent)
    if filepath:
        cmd_data["filepath"] = filepath
    with open(cmd_file, "w") as f:
        json.dump(cmd_data, f)

    # 清理旧结果
    if os.path.exists(result_file):
        os.remove(result_file)

    # 触发宏（利用 LO 单实例机制，转发给已运行的 GUI 实例）
    macro_url = "vnd.sun.star.script:user/office_cmd.py$run_command?language=Python&location=user"
    env = {**os.environ, "DISPLAY": ":0", "WAYLAND_DISPLAY": "wayland-0"}
    proc = subprocess.Popen(
        ["libreoffice", macro_url],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )

    # 等待结果（最多 10 秒）
    for _ in range(20):
        time.sleep(0.5)
        if os.path.exists(result_file):
            try:
                with open(result_file) as f:
                    return json.load(f)
            except Exception:
                pass

    proc.terminate()
    return {"ok": False, "message": "宏执行超时，已降级到 python-docx"}


def exec_writer(intent: dict, filepath: str) -> dict:
    import docx
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    action = intent["action"]
    p = Path(os.path.expanduser(filepath))

    if action == "new_doc":
        doc = docx.Document()
        doc.add_paragraph("新文档")
        doc.save(str(p))
        return {"ok": True, "message": f"已创建 {p}"}

    if action == "get_content":
        doc = docx.Document(str(p))
        lines = [
            f"第{i + 1}段: {para.text}"
            for i, para in enumerate(doc.paragraphs)
            if para.text.strip()
        ]
        return {"ok": True, "message": "\n".join(lines[:30]) or "文档为空"}

    # 修改类操作
    doc = docx.Document(str(p))

    if action == "add_paragraph":
        doc.add_paragraph(intent.get("content", ""))
        doc.save(str(p))
        return {"ok": True, "message": f"已添加段落: {intent.get('content', '')[:30]}"}

    if action == "insert_text":
        last = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
        last.add_run(intent.get("content", ""))
        doc.save(str(p))
        return {"ok": True, "message": f"已插入文字: {intent.get('content', '')[:30]}"}

    if action == "find_replace":
        find = intent.get("find", "")
        replace = intent.get("replace", "")
        count = 0
        for para in doc.paragraphs:
            if find in para.text:
                for run in para.runs:
                    if find in run.text:
                        run.text = run.text.replace(find, replace)
                        count += 1
        doc.save(str(p))
        return {"ok": True, "message": f"已替换 {count} 处: '{find}' → '{replace}'"}

    if action == "insert_table":
        rows = intent.get("rows", 3)
        cols = intent.get("cols", 3)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        doc.save(str(p))
        return {"ok": True, "message": f"已插入 {rows}×{cols} 表格"}

    if action in (
        "bold",
        "italic",
        "underline",
        "font_size",
        "color",
        "set_align",
        "set_line_spacing",
        "set_heading",
        "delete_paragraph",
    ):
        target = intent.get("target", "all")
        target_paras = get_paragraph(doc, target)
        if not target_paras:
            return {
                "ok": False,
                "message": f"段落 {target} 不存在（共{len(doc.paragraphs)}段）",
            }

        for para in target_paras:
            if action == "delete_paragraph":
                p_elem = para._element
                p_elem.getparent().remove(p_elem)
                continue
            if action == "set_heading":
                para.style = f"Heading {intent.get('level', 1)}"
                continue
            if action == "set_align":
                align_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                }
                para.alignment = align_map.get(
                    intent.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT
                )
                continue
            if action == "set_line_spacing":
                from docx.shared import Pt as _Pt

                para.paragraph_format.line_spacing = _Pt(
                    float(intent.get("spacing", 1.5)) * 12
                )
                continue
            for run in para.runs:
                if action == "bold":
                    run.bold = True
                elif action == "italic":
                    run.italic = True
                elif action == "underline":
                    run.underline = True
                elif action == "font_size":
                    run.font.size = Pt(float(intent.get("size", 12)))
                elif action == "color":
                    hex_c = intent.get("color", "000000").lstrip("#")
                    r, g, b = (
                        int(hex_c[0:2], 16),
                        int(hex_c[2:4], 16),
                        int(hex_c[4:6], 16),
                    )
                    run.font.color.rgb = RGBColor(r, g, b)

        doc.save(str(p))
        n = len(target_paras)
        return {"ok": True, "message": f"已对 {n} 个段落执行 {action}"}

    if action == "page_break":
        from docx.enum.text import WD_BREAK

        last_para = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
        run = last_para.add_run()
        run.add_break(WD_BREAK.PAGE)
        doc.save(str(p))
        return {"ok": True, "message": "已插入分页符"}

    if action == "export_pdf":
        pdf_path = intent.get("path") or str(p.with_suffix(".pdf"))
        # 先保存确保最新内容
        doc.save(str(p))
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(Path(pdf_path).parent),
                str(p),
            ],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode == 0:
            return {"ok": True, "message": f"已导出 PDF: {pdf_path}"}
        return {"ok": False, "message": f"PDF导出失败: {result.stderr[:100]}"}

    if action in ("add_comment", "track_changes"):
        # 这两个需要 LO 运行，走宏路径
        if _lo_running():
            return exec_via_macro(intent, filepath)
        return {
            "ok": False,
            "message": f"{action} 需要 LibreOffice 运行中，请先打开文件",
        }

    if action == "save":
        doc.save(str(p))
        return {"ok": True, "message": f"已保存 {p}"}

    if action == "insert_image":
        img_path = intent.get("image_path", "")
        if not img_path or not Path(os.path.expanduser(img_path)).exists():
            return {"ok": False, "message": f"图片不存在: {img_path}"}
        width = intent.get("width")  # cm
        height = intent.get("height")  # cm
        doc.add_picture(
            os.path.expanduser(img_path),
            width=Cm(width) if width else None,
            height=Cm(height) if height else None,
        )
        doc.save(str(p))
        return {"ok": True, "message": f"已插入图片: {img_path}"}

    if action == "header_footer":
        section = doc.sections[0]
        header_text = intent.get("header", "")
        footer_text = intent.get("footer", "")
        if header_text:
            section.header.paragraphs[0].text = header_text
        if footer_text:
            section.footer.paragraphs[0].text = footer_text
        doc.save(str(p))
        parts = []
        if header_text:
            parts.append(f"页眉: {header_text}")
        if footer_text:
            parts.append(f"页脚: {footer_text}")
        return {"ok": True, "message": "已设置 " + ", ".join(parts)}

    return {"ok": False, "message": f"未知操作: {action}"}


# ─── Calc 操作 ────────────────────────────────────────


def exec_calc(intent: dict, filepath: str) -> dict:
    import openpyxl

    action = intent["action"]
    p = Path(os.path.expanduser(filepath))
    wb = openpyxl.load_workbook(str(p))
    sheet_name = intent.get("sheet")
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active

    if action == "read_cell":
        val = ws[intent["cell"]].value
        return {"ok": True, "message": f"{intent['cell']} = {val}"}

    if action == "write_cell":
        ws[intent["cell"]] = intent.get("value")
        wb.save(str(p))
        return {
            "ok": True,
            "message": f"已写入 {intent['cell']} = {intent.get('value')}",
        }

    if action == "write_row":
        row = intent.get("row", 1)
        values = intent.get("values", [])
        for col, val in enumerate(values, 1):
            ws.cell(row=row, column=col, value=val)
        wb.save(str(p))
        return {"ok": True, "message": f"已写入第{row}行: {values}"}

    if action == "get_sheet_summary":
        summary = f"工作表: {ws.title}\n行数: {ws.max_row}\n列数: {ws.max_column}\n"
        for row in ws.iter_rows(max_row=3, values_only=True):
            summary += str(list(row)) + "\n"
        return {"ok": True, "message": summary}

    if action == "read_range":
        from openpyxl.utils import range_boundaries

        rng = intent.get("range", "A1:C5")
        min_col, min_row, max_col, max_row = range_boundaries(rng)
        rows_data = []
        for row in ws.iter_rows(
            min_row=min_row,
            max_row=max_row,
            min_col=min_col,
            max_col=max_col,
            values_only=True,
        ):
            rows_data.append([str(c) if c is not None else "" for c in row])
        return {
            "ok": True,
            "message": f"{rng}:\n" + "\n".join("\t".join(r) for r in rows_data),
            "data": rows_data,
        }

    if action == "find_replace":
        find_txt = intent.get("find", "")
        repl_txt = intent.get("replace", "")
        count = 0
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and find_txt in str(cell.value):
                    cell.value = str(cell.value).replace(find_txt, repl_txt)
                    count += 1
        wb.save(str(p))
        return {
            "ok": True,
            "message": f"已替换 {count} 处: '{find_txt}' → '{repl_txt}'",
        }

    if action == "batch_write":
        cells = intent.get("cells", [])
        count = 0
        for item in cells:
            cell_ref = item.get("cell", "")
            value = item.get("value")
            if cell_ref and value is not None:
                ws[cell_ref] = value
                count += 1
        wb.save(str(p))
        return {"ok": True, "message": f"已批量写入 {count} 个单元格"}

    if action == "delete_row":
        row_num = intent.get("row", 1)
        ws.delete_rows(row_num)
        wb.save(str(p))
        return {"ok": True, "message": f"已删除第 {row_num} 行"}

    if action == "insert_row":
        row_num = intent.get("row", 1)
        values = intent.get("values", [])
        ws.insert_rows(row_num)
        if values:
            for col, val in enumerate(values, 1):
                ws.cell(row=row_num, column=col, value=val)
        wb.save(str(p))
        return {
            "ok": True,
            "message": f"已在第 {row_num} 行插入{len(values)}列"
            if values
            else f"已在第 {row_num} 行插入空行",
        }

    if action == "format_cell":
        cell_ref = intent.get("cell", "A1")
        fmt = intent.get("format", "")
        cell = ws[cell_ref]
        if fmt:
            cell.number_format = fmt
        bold = intent.get("bold")
        if bold is not None:
            cell.font = cell.font.copy(bold=bool(bold))
        color = intent.get("color")
        if color:
            from openpyxl.styles import Color as XlColor, PatternFill

            cell.fill = PatternFill(
                start_color=color.lstrip("#"),
                end_color=color.lstrip("#"),
                fill_type="solid",
            )
        size = intent.get("size")
        if size:
            cell.font = cell.font.copy(size=float(size))
        wb.save(str(p))
        return {"ok": True, "message": f"已格式化 {cell_ref}"}

    if action == "auto_fit":
        for col in ws.columns:
            max_len = 0
            col_letter = col[0].column_letter
            for cell in col:
                if cell.value:
                    max_len = max(max_len, len(str(cell.value)))
            ws.column_dimensions[col_letter].width = min(max_len + 2, 50)
        wb.save(str(p))
        return {"ok": True, "message": "已自动调整列宽"}

    return {"ok": False, "message": f"未知Calc操作: {action}"}


def execute_intent(intent: dict, filepath: Optional[str] = None) -> dict:
    global _current_file
    action = intent.get("action", "unknown")

    # 特殊动作：设置当前文件
    if action == "set_current":
        fp = intent.get("filepath", filepath)
        if fp:
            _current_file = os.path.expanduser(fp)
            return {"ok": True, "message": f"当前文件已设为: {_current_file}"}
        return {"ok": False, "message": "需要指定 filepath"}

    # 打开文件（用LibreOffice GUI显示）
    if action == "open":
        fp = intent.get("filepath", filepath) or _current_file
        if fp:
            _current_file = os.path.expanduser(fp)
        cmd = (
            ["libreoffice", _current_file]
            if _current_file
            else ["libreoffice", "--writer"]
        )
        subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        return {
            "ok": True,
            "message": f"已用 LibreOffice 打开: {_current_file or '新文档'}",
        }

    # 新建文档
    if action == "new_doc":
        fp = os.path.expanduser(intent.get("filepath", "~/Desktop/新文档.docx"))
        _current_file = fp
        return exec_writer(intent, fp)

    # 确定文件路径
    fp = filepath or _current_file
    if not fp and action not in ("open",):
        return {"ok": False, "message": "未设置当前文档。请先说：打开 /路径/文件.docx"}

    if fp:
        _current_file = os.path.expanduser(fp)
        p = Path(_current_file)

        # UNO专属操作，优先走宏（需LO运行）
        UNO_ONLY = {"add_comment", "track_changes"}
        if action in UNO_ONLY and _lo_running():
            return exec_via_macro(intent, _current_file)

        if is_calc(p):
            result = exec_calc(intent, _current_file)
        else:
            result = exec_writer(intent, _current_file)

        # 文件变更通知
        if result.get("ok") and action not in (
            "read_cell",
            "read_sheet",
            "read_document",
        ):
            try:
                fname = Path(_current_file).name if _current_file else "文档"
                subprocess.Popen(
                    ["notify-send", "文档已更新", f"{fname}: {action}"],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
            except Exception:
                pass

        return result

    return {"ok": False, "message": f"未知操作: {action}"}


def save_history(entry: dict):
    with open(HISTORY_FILE, "a") as f:
        f.write(json.dumps(entry, ensure_ascii=False) + "\n")


# ── HTTP API ─────────────────────────────────────────────


@app.get("/health")
async def health():
    return {"ok": True, "current_file": _current_file, "version": "2.1"}


@app.post("/command")
async def command(body: dict):
    """自然语言 → 意图解析 → 执行文档操作"""
    text = body.get("text", "")
    source = body.get("source", "chat")
    filepath = body.get("filepath")
    log.info("[%s] %s", source, text)

    try:
        intent = await parse_intent(text)
        log.info("意图: %s", intent)

        if filepath:
            intent.setdefault("filepath", filepath)

        result = await asyncio.get_event_loop().run_in_executor(
            None, execute_intent, intent, filepath
        )
        save_history(
            {
                "ts": time.time(),
                "source": source,
                "text": text,
                "intent": intent,
                "result": result,
            }
        )

        # 同步到 dialogue feed
        try:
            async with httpx.AsyncClient(timeout=3) as client:
                await client.post(
                    "http://localhost:9800/api/feed-ingest",
                    json={
                        "type": "office",
                        "text": f"[{source}] {text} → {result.get('message', 'done')}",
                    },
                )
        except Exception:
            pass

        return result
    except Exception as e:
        log.error("command error: %s", e, exc_info=True)
        return {"ok": False, "message": str(e)}


@app.post("/execute")
async def execute(body: dict):
    """直接执行已解析的 intent JSON"""
    try:
        result = await asyncio.get_event_loop().run_in_executor(
            None, execute_intent, body, body.get("filepath")
        )
        save_history({"ts": time.time(), "intent": body, "result": result})
        return result
    except Exception as e:
        return {"ok": False, "message": str(e)}


@app.get("/current-file")
async def get_current_file():
    return {"filepath": _current_file}


@app.post("/set-file")
async def set_file(body: dict):
    global _current_file
    _current_file = os.path.expanduser(body.get("filepath", ""))
    return {"ok": True, "filepath": _current_file}


@app.get("/history")
async def history(limit: int = 20):
    if not HISTORY_FILE.exists():
        return {"items": []}
    lines = HISTORY_FILE.read_text().strip().split("\n")
    items = []
    for line in reversed(lines):
        try:
            items.append(json.loads(line))
        except Exception:
            pass
        if len(items) >= limit:
            break
    return {"items": items}


if __name__ == "__main__":
    import sys
    import argparse

    parser = argparse.ArgumentParser(description="Office Agent CLI")
    parser.add_argument(
        "--cli", nargs="?", const="-", help="CLI模式: 直接执行指令，无HTTP服务"
    )
    parser.add_argument("--port", type=int, default=9810, help="HTTP端口")
    args = parser.parse_args()

    if args.cli:
        # CLI 模式：读取指令并执行
        if args.cli == "-":
            text = sys.stdin.read().strip()
        else:
            text = args.cli

        if not text:
            print("[FAIL] 无指令内容", file=sys.stderr)
            sys.exit(1)

        async def cli_run():
            intent = await parse_intent(text)
            log.info("CLI意图: %s", intent)
            result = execute_intent(intent)
            save_history(
                {
                    "ts": time.time(),
                    "source": "cli",
                    "text": text,
                    "intent": intent,
                    "result": result,
                }
            )
            if result.get("ok"):
                print(result["message"])
            else:
                print(f"[FAIL] {result['message']}", file=sys.stderr)
                sys.exit(1)

        asyncio.run(cli_run())
    else:
        uvicorn.run(app, host="0.0.0.0", port=args.port, log_level="info")
